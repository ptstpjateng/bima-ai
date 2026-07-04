"""Official SIAP document templates — fetch, fill, and deliver as editable DOCX.

The PPKP doc-prep co-pilot must hand the citizen the REAL SIAP form (correct
recipient, exact field block), not an invented layout. SIAP hosts these as files
in `ptsp.files` (see services/siap_tools.siap_get_templates); they are served
from Beta storage at `{SIAP_STORAGE_BASE}/storage/{internal_filename}`.

Pipeline per doc_type:
  1. Look up the template's storage key for the resolved licence.
  2. If it is a `.docx` that is actually present on Beta → fetch it and fill the
     citizen's IDENTITY fields (Nama, Jabatan, Alamat, Nomer HP, NIK) in place,
     leaving the vessel fields blank for the citizen. (Decision: identity only.)
  3. Else (legacy `.doc`, or not synced yet, or a fill error) → if BIMA has a
     hand-built DOCX for that doc_type, use it; otherwise return None so the
     caller falls back to the existing fpdf2 generator.

python-docx is imported lazily so this module imports fine where it is absent
(local test env); the container ships it via requirements.txt. Nothing here
raises to the caller — a template miss is a graceful None, never a broken flow.
"""
from __future__ import annotations

import asyncio
import io
import logging
import os
import re
from typing import Any, Optional

import httpx

logger = logging.getLogger("bima_ai.siap_templates")

# Master switch + tunables (env-overridable, no redeploy to flip).
_ENABLED = os.getenv("BIMA_SIAP_TEMPLATES_ENABLED", "true").lower() in ("1", "true", "yes")
_STORAGE_BASE = os.getenv("SIAP_STORAGE_BASE", "https://beta-siap.nolongin.com").rstrip("/")
_FETCH_TIMEOUT = float(os.getenv("BIMA_SIAP_TEMPLATE_TIMEOUT", "20"))
_MAX_TEMPLATE_BYTES = int(os.getenv("BIMA_SIAP_TEMPLATE_MAX_BYTES", str(8 * 1024 * 1024)))
# How many applicant docs one Vision fill pass reads. The vessel/permit specs are
# spread ACROSS documents (Desain Kapal → GT/tipe/bahan/galangan, Spesifikasi →
# alat, Surat Pesanan → thn_bangun, SIUP → no_siup …), so reading a single doc
# leaves most of the form blank. We read up to this many (spec-docs first) in
# parallel and merge. Bounds latency/cost; PKPP has ~8 applicant docs.
_VISION_MAX_DOCS = int(os.getenv("BIMA_VISION_MAX_DOCS", "8"))

_ID_MONTHS = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli",
              "Agustus", "September", "Oktober", "November", "Desember"]


def is_enabled() -> bool:
    return _ENABLED


def _today_id() -> str:
    """Indonesian long date, e.g. '2 Juli 2026'. Runtime-only (not import-time)."""
    from datetime import datetime
    now = datetime.now()
    return f"{now.day} {_ID_MONTHS[now.month]} {now.year}"


# ---------------------------------------------------------------------------
# Identity-field mapping: BIMA session field -> label variants seen in the SIAP
# Word templates. Vessel fields are intentionally NOT here (left blank for the
# citizen). Matching is case-insensitive on the label text before the ':'.
# ---------------------------------------------------------------------------
_IDENTITY_LABELS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("applicant_name", ("nama pemohon", "nama lengkap", "nama")),
    ("jabatan", ("jabatan",)),
    ("alamat", ("alamat",)),
    ("phone", ("nomer hp", "nomor hp", "no hp", "no. hp", "no telp", "no. telp",
               "nomor telepon", "telepon", "hp")),
    ("nik", ("nik",)),
)


def _identity_value(field_key: str, fields: dict[str, Any]) -> Optional[str]:
    """Resolve the value BIMA has for an identity field, or None to leave blank.

    NO-HALLUCINATION CONTRACT (M2): every field, including `jabatan`, returns the
    citizen's REAL value or None — no hardcoded default. `jabatan` used to default
    to "Pemilik Kapal", which stamped a guessed occupation onto the citizen's own
    Surat Permohonan; that is a fabricated value and is removed. An unset jabatan
    now renders blank for the citizen to fill (the standalone "Pemilik Kapal"
    caption below the signature block in build_surat_permohonan_ppkp_docx is a
    fixed template caption, not a per-citizen data fill, and is left as-is)."""
    val = fields.get(field_key)
    if val is None:
        return None
    val = str(val).strip()
    return val or None


def _match_identity_field(label_text: str) -> Optional[str]:
    """Return the field_key whose label matches this line's pre-colon text."""
    low = label_text.strip().lower()
    for field_key, variants in _IDENTITY_LABELS:
        variants_t = (variants,) if isinstance(variants, str) else variants
        for v in variants_t:
            # Whole-label match (allow a trailing '*' or spaces), not substring,
            # so "Nama" does not swallow "Nama Kapal".
            if low == v or low.rstrip(" *") == v:
                return field_key
    return None


# ---------------------------------------------------------------------------
# Fetch
# ---------------------------------------------------------------------------
async def fetch_template_bytes(internal_filename: str) -> Optional[bytes]:
    """GET the template file from Beta SIAP storage. None on any miss/error.

    Not-yet-synced files return 404 → None → caller falls back. Never raises.
    """
    # internal_filename comes from the SIAP DB (ULID paths); guard anyway so a
    # bad/tampered value can't traverse out of /storage or hit another host.
    if (not internal_filename or ".." in internal_filename
            or internal_filename.startswith(("/", "\\"))):
        if internal_filename:
            logger.warning("SIAP template ref rejected (unsafe path)")
        return None
    url = f"{_STORAGE_BASE}/storage/{internal_filename}"
    try:
        # A public static file returns 200 directly; a 30x is anomalous and could
        # redirect off the trusted storage host (SSRF), so refuse to follow it.
        async with httpx.AsyncClient(timeout=_FETCH_TIMEOUT, follow_redirects=False) as client:
            resp = await client.get(url)
        if resp.status_code != 200:
            logger.info("SIAP template fetch miss | status=%d file=%s",
                        resp.status_code, internal_filename)
            return None
        data = resp.content
        if not data or len(data) > _MAX_TEMPLATE_BYTES:
            logger.warning("SIAP template size out of range | bytes=%d file=%s",
                           len(data or b""), internal_filename)
            return None
        return data
    except Exception:
        logger.exception("SIAP template fetch failed | file=%s", internal_filename)
        return None


# ---------------------------------------------------------------------------
# Submission-document byte fetch (officer copilot — [[Multi-step Officer Copilot]])
#
# The next-step officer must be able to open the citizen's uploaded requirement
# files STRAIGHT FROM SIAP. Those files live in the same Beta storage as the
# blank templates, but under `storage/app/public/` (Laravel's public disk) with
# the DB storing a relative key (`profile_requirements.profile_requirements_files`).
# We reuse the exact same hardened GET (SSRF-safe: no redirects, size-capped,
# never raises, never logs bytes) as `fetch_template_bytes`, only differing in
# the storage sub-prefix. The sub-prefix is env-overridable because SIAP's
# public-disk layout is environment-specific; the DB value may already include
# the `app/public/` segment, so we try the value verbatim first, then prefixed.
# ---------------------------------------------------------------------------

# Laravel public-disk prefix under `/storage/`. The stored key is usually a bare
# path like "berkas/2026/xyz.pdf"; the public URL is
# {BASE}/storage/app/public/<key>. Overridable per environment; empty → the
# key is fetched under `/storage/<key>` verbatim (same as templates).
_SUBMISSION_STORAGE_PREFIX = os.getenv(
    "SIAP_SUBMISSION_STORAGE_PREFIX", "app/public"
).strip("/")


def _submission_candidate_paths(file_ref: str) -> list[str]:
    """Ordered storage paths to try for one submission-file key.

    A DB value may already carry its own storage sub-path, so we try it
    verbatim first, then under the public-disk prefix. De-duplicated; unsafe
    keys are filtered by fetch_template_bytes itself."""
    ref = (file_ref or "").strip().lstrip("/")
    if not ref:
        return []
    candidates = [ref]
    if _SUBMISSION_STORAGE_PREFIX and not ref.startswith(_SUBMISSION_STORAGE_PREFIX + "/"):
        candidates.append(f"{_SUBMISSION_STORAGE_PREFIX}/{ref}")
    # Preserve order, drop dupes.
    seen: set[str] = set()
    out: list[str] = []
    for c in candidates:
        if c not in seen:
            seen.add(c)
            out.append(c)
    return out


async def fetch_submission_file_bytes(file_ref: str) -> Optional[bytes]:
    """GET a citizen-uploaded requirement file from Beta SIAP storage.

    `file_ref` is `ptsp.profile_requirements.profile_requirements_files`. Tries
    the value verbatim, then under the public-disk prefix; returns the first
    200-with-bytes hit. None on any miss/error. Never raises; never logs the
    file name or the bytes (both may echo applicant PII)."""
    for candidate in _submission_candidate_paths(file_ref):
        data = await fetch_template_bytes(candidate)
        if data:
            return data
    return None


# ---------------------------------------------------------------------------
# Fill an existing .docx template in place (identity fields only)
# ---------------------------------------------------------------------------
def _iter_block_paragraphs(container):
    """Yield every paragraph in a doc OR table cell, recursing into nested tables."""
    for p in container.paragraphs:
        yield p
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_block_paragraphs(cell)


def _set_paragraph_value_after_colon(paragraph, value: str) -> bool:
    """If `paragraph` is a 'Label : <blank>' line, write value after the colon
    without disturbing the label's formatting. Returns True if it filled."""
    text = paragraph.text
    if ":" not in text:
        return False
    label, _, rest = text.partition(":")
    if _match_identity_field(label) is None:
        return False
    # Only fill when the value slot is empty/placeholder (dots/underscores/dashes).
    if re.sub(r"[\s._–\-]", "", rest):
        return False
    # Append the value to the run that carries the colon (preserves font/size).
    for run in reversed(paragraph.runs):
        if ":" in run.text:
            run.text = run.text[: run.text.rindex(":") + 1] + f" {value}"
            return True
    # Fallback: append a new run at the paragraph end.
    paragraph.add_run(f" {value}")
    return True


def fill_docx_identity(docx_bytes: bytes, fields: dict[str, Any]) -> bytes:
    """Fill identity fields into a .docx template, return the edited .docx bytes.

    Best-effort per label; unmatched fields are left as the citizen must fill.
    NOTE: label variants live in _IDENTITY_LABELS — validate against the real
    SIAP templates and extend there if a form uses a different wording.
    """
    from docx import Document  # lazy: only needed when actually filling

    doc = Document(io.BytesIO(docx_bytes))
    filled: list[str] = []
    for paragraph in _iter_block_paragraphs(doc):
        text = paragraph.text
        if ":" not in text:
            continue
        label = text.split(":", 1)[0]
        field_key = _match_identity_field(label)
        if field_key is None:
            continue
        value = _identity_value(field_key, fields)
        if not value:
            continue
        if _set_paragraph_value_after_colon(paragraph, value):
            filled.append(field_key)
    logger.info("SIAP template filled | identity_fields=%s", ",".join(filled) or "none")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------------------------------------------------------------------------
# Placeholder-token fill — the OFFICER-facing SK / Persetujuan output template.
#
# Unlike the citizen forms ("Label : <blank>"), SIAP's LICENSE-RECOMMEND output
# template (e.g. "Surat PKPP V2.docx") fills via literal `[data.KEY]` tokens
# embedded in the prose. Word almost always SPLITS such a token across several
# runs inside a paragraph (spell-check / formatting boundaries) — in the real
# PKPP template only 2 of 16 tokens survive intact in a single run; the other 14
# are shattered (e.g. "[", "data.no_siup", "]"). A naive per-run replace would
# never see them. So we operate on the JOINED paragraph/cell text, substitute
# every token, and rewrite the paragraph in place — collapsing its runs into one
# so the replacement is whole. Formatting of the paragraph's first run is
# preserved; intra-paragraph run styling is intentionally flattened (these are
# body-copy fill lines, not styled tables).
#
# Unmapped or empty keys are replaced with a blank fill line "………………" (never the
# raw "[data.KEY]", which reads as broken to the officer) so the officer can
# complete them by hand.
# ---------------------------------------------------------------------------

# The literal fill tokens BIMA understands. TWO syntaxes coexist on Beta,
# cleanly split by stereotype:
#   * LICENSE-RECOMMEND (Rekomtek) templates use  [data.KEY]
#   * LICENSE-SK        templates use PhpWord      ${key}
# Keys are snake_case but some templates carry a trailing digit (catatan1,
# syarat1, pernyataan_1), so allow digits — an [A-Za-z_]+ class silently dropped
# those. Both tokens are read off the JOINED paragraph text (run-aware) and both
# route through the SAME classifier/resolver, so a template using either syntax
# fills identically and an unresolved key always becomes the blank fill line.
_PLACEHOLDER_RE = re.compile(r"\[data\.([A-Za-z0-9_]+)\]")
_PHPWORD_RE = re.compile(r"\$\{([A-Za-z0-9_]+)\}")

# What an unmapped/blank field becomes — a visible fill line, not a broken token.
_BLANK_FILL = "…" * 12


def _fill_placeholders_in_text(text: str, data: dict[str, Any]) -> tuple[str, list[str], list[str]]:
    """Replace every `[data.KEY]` AND every `${key}` in `text`. Returns
    (new_text, filled_keys, blank_keys). A key present in `data` with a non-empty
    value fills; anything else (unknown key OR empty value) becomes the blank
    fill line — so an SK template's `${qrcode_signed}` / `${ttd_gbr}` image
    anchors (classified BLANK) render as a fill line, never a raw token."""
    filled: list[str] = []
    blanks: list[str] = []

    def _sub(m: "re.Match[str]") -> str:
        key = m.group(1)
        raw = data.get(key)
        val = "" if raw is None else str(raw).strip()
        if val:
            filled.append(key)
            return val
        blanks.append(key)
        return _BLANK_FILL

    text = _PLACEHOLDER_RE.sub(_sub, text)
    text = _PHPWORD_RE.sub(_sub, text)
    return text, filled, blanks


def _rewrite_paragraph_text(paragraph, new_text: str) -> None:
    """Set a paragraph's full text to `new_text`, preserving the first run's
    formatting and dropping the rest. python-docx has no whole-paragraph setter
    that keeps styling, so we write into run[0] and blank the tail runs (leaving
    them in place keeps the paragraph's XML valid)."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def fill_docx_placeholders(docx_bytes: bytes, data: dict[str, Any]) -> bytes:
    """Fill an SIAP LICENSE-RECOMMEND output template that uses `[data.KEY]`
    tokens, returning the edited .docx bytes.

    Run-aware: a token split across runs is still matched because we join the
    paragraph/cell text before substituting (test: a "[" / "data.x" / "]" run
    split fills correctly). Unmapped/empty keys become a blank fill line so the
    officer completes them; the raw token is never left behind.
    """
    from docx import Document  # lazy: only needed when actually filling

    doc = Document(io.BytesIO(docx_bytes))
    all_filled: list[str] = []
    all_blank: list[str] = []
    for paragraph in _iter_block_paragraphs(doc):
        text = paragraph.text
        if "[data." not in text and "${" not in text:
            continue
        new_text, filled, blanks = _fill_placeholders_in_text(text, data)
        if new_text != text:
            _rewrite_paragraph_text(paragraph, new_text)
            all_filled.extend(filled)
            all_blank.extend(blanks)

    logger.info(
        "SIAP output template filled | fields=%d blank=%d",
        len(set(all_filled)), len(set(all_blank)),
    )
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ===========================================================================
# GENERIC output-template fill engine
# ([[Generalize copilot SK/output-doc from real data]] — task #80.)
#
# The PKPP path above (build_sk_data + _SK_*_KEYS) hardcodes ONE licence's 16
# `[data.KEY]` keys. Every SIAP licence has a DIFFERENT placeholder SET (11–36
# keys observed). So this engine drives the fill DATA-FIRST:
#
#   1. DISCOVER the template's `[data.KEY]` keys dynamically from the real
#      fetched docx (run-aware — same JOINED-paragraph read fill_docx_placeholders
#      uses, so a token shattered across runs is still found).
#   2. RESOLVE each discovered key through a data-source resolver: a key-name
#      classifier routes it to the REAL person_profile.properties, the real
#      license_request/license/step case meta, ONE best-effort Vision pass over
#      the submission docs, or LEAVES IT BLANK.
#   3. Feed the resolved {key: value} map to `fill_docx_placeholders`, which
#      renders any absent/empty key as the `………………` blank fill line.
#
# THE NO-HALLUCINATION CONTRACT (structural, not prose):
#   The resolver ONLY ever puts a key into the data map when it read a value
#   from a real source (a profile column, a case DB row, or text literally on a
#   submitted document). A key it cannot ground is simply NOT added → the
#   renderer emits it as a blank fill line the officer completes by hand. There
#   is no branch that fabricates, defaults, or guesses a value. SIAP-issued
#   values (SK number, penetapan/TTE date) are classified BLANK on purpose.
# ===========================================================================


def discover_placeholder_keys(docx_bytes: bytes) -> list[str]:
    """Return the ordered, de-duplicated set of `[data.KEY]` keys present in a
    template docx — run-aware (a token split across runs is still found because
    we read the JOINED paragraph/cell text, exactly like fill_docx_placeholders).

    Never raises to the caller's flow beyond a python-docx import error (which
    the orchestrator guards). Returns [] for a template with no tokens."""
    from docx import Document  # lazy: only needed when actually discovering

    doc = Document(io.BytesIO(docx_bytes))
    seen: set[str] = set()
    keys: list[str] = []
    for paragraph in _iter_block_paragraphs(doc):
        text = paragraph.text
        if "[data." not in text and "${" not in text:
            continue
        # Both token syntaxes ([data.KEY] and PhpWord ${key}); read off the
        # joined paragraph text so a run-shattered token is still found.
        for regex in (_PLACEHOLDER_RE, _PHPWORD_RE):
            for m in regex.finditer(text):
                key = m.group(1)
                if key not in seen:
                    seen.add(key)
                    keys.append(key)
    return keys


# --- Source classification: which real source a placeholder key draws from ---
#
# Four classes:
#   PROFILE  — an APPLICANT-identity attribute → person_profile.properties.
#   CASE     — a request/licence/step attribute → get_request_case_meta / DB.
#   VISION   — a value that lives ONLY inside an uploaded document.
#   BLANK    — SIAP-issued (SK/registration number, penetapan/TTE date) OR
#              unresolvable → always the blank fill line, never sourced.
#
# The classifier is a data-driven key→profile-field map with sensible key-name
# heuristics, so a NEW licence template works with NO code change: a key that
# names an applicant attribute resolves from the profile; a key that names a
# case attribute from the case meta; a SIAP-issued key stays blank; anything
# else falls to Vision (best-effort) and, failing that, blank.

SOURCE_PROFILE = "profile"
SOURCE_CASE = "case"
SOURCE_VISION = "vision"
SOURCE_BLANK = "blank"
# A key whose value came straight from the request's filled SIAP Formulir Isian
# (ptsp.form_value.properties). This is REAL data — citizen-/officer-entered or
# BIMA-auto-filled from real reads — so it preserves the no-hallucination
# contract. When an `overrides` map (the form_value) carries a NON-BLANK value
# for a key, that value WINS over every other source and no Vision/profile/case
# read is done for that key. A blank/whitespace override is IGNORED (it never
# overrides a real source with emptiness, and never fabricates).
SOURCE_FORM = "form"

# Applicant-identity placeholder-key → person_profile.properties field.
# Covers the exact key names seen across the investigated templates plus common
# synonyms. The person_profile key set is siap_profile_client._OPTIONAL_FIELDS
# + full_name + identity_number (see that module). Every mapping is to a REAL
# profile column — no derived/guessed values.
_PROFILE_KEY_MAP: dict[str, str] = {
    # name
    "nama_pemohon": "full_name",
    "nama": "full_name",
    "full_name": "full_name",
    "atas_nama": "full_name",
    # NB: atas_nama_perusahaan / alamat_perusahaan are NOT here — a COMPANY is
    # not the applicant person, so they route to BLANK (see C2 exclusion in
    # classify_placeholder_key). Mapping them to the person profile would stamp
    # the human's name/address where the company block belongs.
    # address
    "alamat": "address",
    "alamat_pemohon": "address",
    "address": "address",
    # NIK / identity number
    "nik": "identity_number",
    "no_ktp": "identity_number",
    "nomor_ktp": "identity_number",
    "identity_number": "identity_number",
    # phone
    "no_hp": "mobile_phone",
    "nomor_hp": "mobile_phone",
    "hp": "mobile_phone",
    "telepon": "phone",
    "telp": "phone",
    "no_telp": "phone",
    "nomor_telepon": "phone",
    "phone": "phone",
    "mobile_phone": "mobile_phone",
    # email
    "email": "email",
    # occupation / position
    "pekerjaan_pemohon": "position_name",
    "jabatan": "position_name",
    "pekerjaan": "position_name",
    "position_name": "position_name",
    # nip (for officials-as-applicants)
    "nip": "nip",
    # birth
    "tempat_lahir": "birth_place",
    "tgl_lahir": "birth_date",
    "tanggal_lahir": "birth_date",
    # locality
    "kabupaten": "regency",
    "kab_kota": "regency",
    "kecamatan": "sub_district",
    "kelurahan": "village",
    "desa": "village",
}

# Substring heuristics for PROFILE keys the exact map misses (a new template
# that names a field slightly differently). Ordered longest-first so a more
# specific hint wins. Each maps a key-name substring → profile field.
#
# ⚠️ These substrings (jabatan, alamat, pekerjaan) also occur inside keys that
# name an OFFICIAL / OFFICE / COMPANY (jabatan_pejabat, alamat_balai,
# alamat_perusahaan). Those are NOT the applicant, so `_names_official_or_company`
# is consulted BEFORE these hints ever fire — see classify_placeholder_key.
_PROFILE_KEY_HINTS: tuple[tuple[str, str], ...] = (
    ("nama_pemohon", "full_name"),
    ("alamat", "address"),
    ("pekerjaan", "position_name"),
    ("jabatan", "position_name"),
    ("email", "email"),
)

# C2 — official / office / company markers. A key that names the SIGNING
# OFFICIAL, the ISSUING OFFICE, or the applicant's COMPANY is NOT the person
# profile: it must never be filled from person_profile.properties (which holds
# the human applicant). Any placeholder key containing one of these substrings
# is routed to BLANK before any PROFILE match. The applicant is a natural
# person; "atas nama perusahaan" / "alamat balai" / "jabatan pejabat" are a
# different subject entirely and are officer/office-authored, so they stay blank
# fill lines. `pemohon` is deliberately absent — nama_pemohon / alamat_pemohon
# ARE the applicant and must stay PROFILE.
_OFFICIAL_OR_COMPANY_MARKERS: tuple[str, ...] = (
    "pejabat", "penandatangan", "kepala", "balai", "dinas", "kantor",
    "instansi", "perusahaan", "badan_usaha", "plt", "cabdin", "skpd",
)

# Case-derived placeholder key → the get_request_case_meta field (or a
# well-known derivation). Values come from real license_request / license / step
# rows. `jenis_pengadaan` is derived deterministically from the licence NAME.
# Only keys whose value is a real case attribute belong here.
_CASE_KEY_TICKET = ("no_tiket", "nomor_tiket", "ticket", "ticket_num", "no_tiket_permohonan")
_CASE_KEY_LICENSE_NAME = ("nama_izin", "nama_perizinan", "jenis_izin", "nama_bidang", "bidang_izin")
_CASE_KEY_JENIS_PENGADAAN = ("jenis_pengadaan",)

# SIAP-ISSUED placeholder keys → ALWAYS blank (class d). These are only knowable
# at issuance (the SK/registration number, the penetapan/TTE date, the QR/
# barcode, the signing official's block). NEVER sourced. Matched by exact key OR
# a substring hint so a new template's issuance field is also caught.
_BLANK_KEY_EXACT = frozenset({
    "no_ppkp", "no_sk", "nomor_sk", "no_rekom", "nomor_rekom", "no_pertek",
    "nomor_rekomtek", "no_persetujuan", "rekom_tgl_penetapan", "tgl_penetapan",
    "mulai_tanggal_sk", "sampai_tanggal_sk", "masa_berlaku", "tgl_berlaku",
    "tgl_hbs_berlaku", "qrcode_signed", "ttd_gbr", "no_reg", "nomor_reg",
})

# SIAP-ISSUED / signature substring hints (C1c). Broadened SURGICALLY to catch
# the issuance number/date/signature/validity synonyms the exact set misses —
# WITHOUT blanket-blanking the bare `no_`/`tgl_` prefixes (which would wrongly
# kill no_siup / tgl_siup — the applicant's prior SIUP, physically on an uploaded
# doc — and tgl_surat_permohonan — the readable application-letter date). Every
# hint here is anchored to an ISSUANCE NOUN, a SIGNATURE noun, or a VALIDITY
# noun, never a bare number/date prefix.
_BLANK_KEY_HINTS: tuple[str, ...] = (
    # issuance / registration number & date nouns
    "penetapan", "ditetapkan", "diterbitkan", "terbit",
    "reg", "registrasi", "agenda", "verifikasi",
    # the SK / decision / recommendation nouns themselves — a no_/tgl_/nomor_
    # bound to one of these is SIAP-issued (no_sk, tgl_sk, nomor_keputusan…)
    "sk", "keputusan", "rekom", "pertek", "persetujuan",
    # QR / barcode / verification anchors
    "qrcode", "qr_code", "barcode", "kode_verifikasi",
    # signature / signing-official block
    "ttd_", "tanda_tangan", "penandatangan", "penandatanganan",
    "pejabat", "tmt", "_signed",
    # validity window
    "masa_", "berlaku", "mulai_tanggal_sk", "sampai_tanggal_sk",
    # officer-authored SK clause / free-text fields (H2) — the officer writes
    # these by hand; NEVER Vision-fabricated clause text.
    "catatan", "syarat", "pernyataan", "keterangan", "isian",
)

# VISION allow-list (C1b). A key qualifies for a Gemini Vision OCR pass ONLY if
# it names a value that PHYSICALLY LIVES ON a citizen-uploaded document. Built
# from the real templates seen (vessel / land / prior-permit). Anything that is
# not PROFILE, not CASE, and not on this allow-list falls through to BLANK — the
# inverted default that makes an unrecognized issuance-field synonym render as a
# blank fill line rather than an OCR-fabricated value stamped as official data.
_VISION_KEY_HINTS: tuple[str, ...] = (
    # vessel (PPKP) — data off the vessel spec / surat permohonan
    "kapal", "gt", "alat", "bahan", "galangan", "tipe",
    "thn_bangun", "tahun_bangun",
    # land (pertanahan / pengairan) — off the land situation drawing / survey
    "tanah", "luas", "batas", "status_tanah", "dimanfaatkan", "peruntukan",
    "maksud", "tujuan", "lokasi", "desa_lokasi", "kab_lokasi", "kec_lokasi",
    "situasi", "gbr_situasi",
    # applicant PRIOR-PERMIT fields that ARE on an uploaded doc (the citizen's
    # existing SIUP). Kept fillable via Vision, and explicitly protected from the
    # `no_`/`tgl_` blank sweep by anchoring the blank hints to issuance nouns.
    "siup",
    # the APPLICATION LETTER (surat permohonan) — number/date/subject are read
    # off the citizen's own letter, not SIAP-issued. tgl_surat_permohonan /
    # tgl_srt_permohonan / no_surat_perm / perihal_surat_perm stay fillable. This
    # is the applicant's letter, NOT the SK, so it is a Vision read, not blank.
    "permohonan", "surat_perm", "srt_perm",
)


def _names_official_or_company(k: str) -> bool:
    """True if the key names the SIGNING OFFICIAL, the ISSUING OFFICE, or the
    applicant's COMPANY — i.e. NOT the applicant person. Such a key must route to
    BLANK before any PROFILE substring (jabatan/alamat) can fire on it (C2)."""
    return any(m in k for m in _OFFICIAL_OR_COMPANY_MARKERS)


def classify_placeholder_key(key: str) -> str:
    """Route a `[data.KEY]` key to its real data source class.

    Returns one of SOURCE_PROFILE / SOURCE_CASE / SOURCE_VISION / SOURCE_BLANK.

    THE INVERTED DEFAULT (no-hallucination core): an unrecognized key is BLANK,
    NOT Vision. A value goes to Vision ONLY when it matches a known applicant-
    DOCUMENT data pattern (the _VISION_KEY_HINTS allow-list). This kills the old
    hole where an issuance-field synonym on a new SK template (nomor_surat,
    tgl_ditetapkan, kode_verifikasi…) fell through to Vision and got a number/
    date OCR'd off some uploaded doc and stamped as the SK's official issuance
    data.

    Resolution ladder (STOP at first hit):
      1. SIAP-issued / signature / validity / officer-clause → BLANK.
      2. official / office / company key → BLANK (never the applicant person).
      3. applicant identity (exact map, then substring hint) → PROFILE.
      4. case attribute (ticket / licence name / jenis_pengadaan) → CASE.
      5. applicant-DOCUMENT data pattern (allow-list) → VISION (best-effort;
         a Vision miss renders blank anyway).
      6. else → BLANK (the inverted default — never a guessed/OCR value)."""
    k = (key or "").strip().lower()
    if not k:
        return SOURCE_BLANK
    # 1) SIAP-issued / signature / validity / officer-clause → never sourced.
    if k in _BLANK_KEY_EXACT or any(h in k for h in _BLANK_KEY_HINTS):
        return SOURCE_BLANK
    # 2) official / office / company → BLANK (NOT the applicant profile). This
    #    runs BEFORE the profile match so jabatan_pejabat / alamat_balai /
    #    atas_nama_perusahaan never draw the human applicant's data.
    if _names_official_or_company(k):
        return SOURCE_BLANK
    # 3) applicant identity → profile.
    if k in _PROFILE_KEY_MAP:
        return SOURCE_PROFILE
    for hint, _field in _PROFILE_KEY_HINTS:
        if hint in k:
            return SOURCE_PROFILE
    # 4) case attribute → case meta / DB.
    if (k in _CASE_KEY_TICKET or k in _CASE_KEY_LICENSE_NAME
            or k in _CASE_KEY_JENIS_PENGADAAN):
        return SOURCE_CASE
    # 5) applicant-DOCUMENT data pattern → Vision (allow-list only).
    if any(h in k for h in _VISION_KEY_HINTS):
        return SOURCE_VISION
    # 6) INVERTED DEFAULT — unrecognized → BLANK, never a guessed/OCR value.
    return SOURCE_BLANK


def _profile_field_for_key(key: str) -> Optional[str]:
    """The person_profile.properties field a PROFILE key draws from, or None."""
    k = (key or "").strip().lower()
    if k in _PROFILE_KEY_MAP:
        return _PROFILE_KEY_MAP[k]
    for hint, field in _PROFILE_KEY_HINTS:
        if hint in k:
            return field
    return None


def _resolve_profile_value(key: str, profile: dict) -> str:
    """Read a PROFILE placeholder value from the real person_profile.properties.
    Returns '' when the field is absent/empty (→ blank fill line). Never guesses.
    """
    if not profile or not isinstance(profile, dict):
        return ""
    field = _profile_field_for_key(key)
    if not field:
        return ""
    val = profile.get(field)
    if val is None:
        return ""
    return str(val).strip()


def _resolve_case_value(key: str, case: Optional[dict], license_name: Optional[str]) -> str:
    """Read a CASE placeholder value from the real case meta. Returns '' when it
    cannot be sourced (→ blank). `case` is the get_request_case_meta dict;
    `license_name` is the resolved ptsp.license.name."""
    k = (key or "").strip().lower()
    case = case or {}
    if k in _CASE_KEY_TICKET:
        return str(case.get("ticket") or "").strip()
    if k in _CASE_KEY_LICENSE_NAME:
        return str(license_name or "").strip()
    if k in _CASE_KEY_JENIS_PENGADAAN:
        return _jenis_pengadaan_from_license(license_name)
    return ""


# Generic Vision extraction: for VISION-classified keys we ask Gemini to read
# ONLY the requested keys off the best submission doc, returning "" for anything
# not literally visible. The schema is BUILT DYNAMICALLY from the discovered
# keys, so there is no per-licence hardcode — a new template's doc-only fields
# are extracted by their own key names. Human-ish descriptions help the model;
# the key name itself is the contract.
_VISION_KEY_DESCRIPTIONS: dict[str, str] = {
    "nama_kapal": "Nama kapal perikanan.",
    "gt": "Ukuran/tonase kapal (Gross Tonnage), angka saja bila ada.",
    "bahan": "Bahan kasko kapal (Kayu/Fiberglass/Besi/Baja).",
    "thn_bangun": "Tahun pembuatan/pembangunan kapal.",
    "alat": "Alat penangkap ikan yang digunakan.",
    "galangan": "Nama/alamat galangan atau tukang pembuat kapal.",
    "no_siup": "Nomor SIUP (Surat Izin Usaha Perikanan).",
    "tgl_siup": "Tanggal terbit SIUP.",
    "tgl_srt_permohonan": "Tanggal surat permohonan pemohon.",
    "tgl_surat_permohonan": "Tanggal surat permohonan pemohon.",
    "tgl_surat_perm": "Tanggal surat permohonan pemohon.",
    "no_surat_perm": "Nomor surat permohonan pemohon.",
    "nomor_surat_permohonan": "Nomor surat permohonan pemohon.",
    "tipe": "Tipe/jenis kapal.",
    "perihal_surat_perm": "Perihal/subjek surat permohonan.",
    "perihal_surat_perm_": "Perihal/subjek surat permohonan.",
    "luas_tanah": "Luas tanah/lahan (dalam m2 bila ada).",
    "status_tanah": "Status kepemilikan tanah.",
    "jenis_tanah": "Jenis tanah/peruntukan.",
    "batas_tanah": "Batas-batas tanah.",
    "no_siup_": "Nomor SIUP.",
    "peruntukan": "Peruntukan/penggunaan yang dimohonkan.",
    "no_pertek_perm": "Nomor pertek pada surat permohonan.",
}


def _vision_description_for_key(key: str) -> str:
    """A short Indonesian description to guide Vision for one key. Falls back to
    a generic 'read this field off the document, empty if not visible' line so
    an UNKNOWN doc-only key on a new template is still extracted by its name."""
    k = (key or "").strip().lower()
    if k in _VISION_KEY_DESCRIPTIONS:
        return _VISION_KEY_DESCRIPTIONS[k]
    pretty = k.replace("_", " ")
    return (
        f"Nilai '{pretty}' persis seperti tertulis pada dokumen. "
        "Kosongkan bila tidak terlihat."
    )


_GENERIC_VISION_PROMPT = (
    "Anda asisten petugas perizinan. Dari dokumen permohonan terlampir (surat "
    "permohonan, izin usaha, spesifikasi/berkas teknis), ambil HANYA data yang "
    "benar-benar TERTULIS di dokumen untuk setiap field yang diminta. Untuk "
    "field yang tidak terlihat, kembalikan string kosong. JANGAN mengarang — "
    "hanya yang terbaca."
)


async def _extract_vision_fields(
    vision_keys: list[str], documents: Optional[dict]
) -> dict[str, str]:
    """Run a best-effort Gemini Vision pass over the submission docs to pull the
    VISION-classified keys. Reads up to `_VISION_MAX_DOCS` applicant docs
    (spec-bearing first) concurrently and MERGES their reads — the vessel/permit
    specs live in DIFFERENT documents, so a single-doc pass leaves most fields
    blank. Schema is built from `vision_keys` so it is licence-agnostic. Returns
    {key: value} for the non-empty reads only, or {} when Vision is unconfigured
    / fails / there are no bytes / no vision keys. NEVER raises. Bytes never
    logged.

    `documents` is the copilot `_doc_context` shape:
      {file_id: {"filename","mime_type","content","claimed_type","detected_type"}}
    """
    if not vision_keys or not documents or not isinstance(documents, dict):
        return {}
    try:
        from services.gemini_vision import extract_structured, is_configured
    except Exception:  # pragma: no cover — vision module import guard
        return {}
    if not is_configured():
        logger.info("generic Vision extract: Gemini Vision not configured")
        return {}

    docs_with_bytes = [d for d in documents.values() if d.get("content")]
    if not docs_with_bytes:
        return {}

    def _label(d: dict) -> str:
        return " ".join(
            str(d.get(k) or "") for k in ("detected_type", "claimed_type", "filename")
        ).lower()

    # Read EACH applicant doc and MERGE — the vessel/permit specs are spread
    # across documents (GT/tipe/bahan/galangan in the Desain Kapal, alat in the
    # Spesifikasi, thn_bangun in the Surat Pesanan, no_siup/tgl_siup in the SIUP,
    # dates/perihal in the Surat Permohonan). A single-doc pass left most of the
    # form blank (exactly what the officer saw). Spec-bearing docs go first so
    # their value wins a per-key tie; still NO hallucination — each pass extracts
    # only what is literally in that one doc, and merging real reads stays real.
    _spec_words = (
        "desain", "rancang", "gambar", "spesifikasi", "teknis", "kapal",
        "pesanan", "kontrak", "permohonan", "siup", "izin usaha", "nib",
    )
    ordered = sorted(
        docs_with_bytes,
        key=lambda d: 0 if any(w in _label(d) for w in _spec_words) else 1,
    )[:_VISION_MAX_DOCS]

    schema: dict[str, Any] = {
        "type": "object",
        "properties": {
            k: {"type": "string", "description": _vision_description_for_key(k)}
            for k in vision_keys
        },
        "required": [],
    }

    async def _extract_one(d: dict) -> Optional[dict]:
        try:
            return await extract_structured(
                image_bytes=d["content"],
                mime_type=d.get("mime_type", "application/octet-stream"),
                prompt=_GENERIC_VISION_PROMPT,
                response_schema=schema,
            )
        except Exception:  # pragma: no cover — extract_structured already guards
            logger.exception("generic Vision extract raised on one doc (skipped)")
            return None

    # Run the per-doc passes concurrently (wall-clock ≈ one pass); gather keeps
    # `ordered` order so the spec-doc-first, first-non-blank-wins merge is stable.
    results = await asyncio.gather(*[_extract_one(d) for d in ordered])

    out: dict[str, str] = {}
    for parsed in results:
        if not parsed or not isinstance(parsed, dict):
            continue
        for key in vision_keys:
            if key in out:  # first non-blank read (spec-doc order) wins
                continue
            val = str(parsed.get(key) or "").strip()
            if val:
                out[key] = val
    logger.info(
        "generic Vision extract | docs=%d requested=%d filled=%d",
        len(ordered), len(vision_keys), len(out),
    )
    return out


def _override_value(key: str, overrides: Optional[dict]) -> str:
    """Return the NON-BLANK override value for a key from the SIAP form_value
    map, or '' when there is none. A blank/whitespace value is IGNORED (returns
    '') so an empty form field never overrides a real source with emptiness and
    never fabricates. Non-string values are stringified then trimmed."""
    if not overrides or not isinstance(overrides, dict):
        return ""
    if key not in overrides:
        return ""
    raw = overrides.get(key)
    if raw is None:
        return ""
    return str(raw).strip()


async def resolve_fill_data(
    keys: list[str],
    *,
    profile: Optional[dict] = None,
    case_meta: Optional[dict] = None,
    license_name: Optional[str] = None,
    documents: Optional[dict] = None,
    run_vision: bool = True,
    overrides: Optional[dict] = None,
) -> tuple[dict[str, str], dict[str, str]]:
    """Resolve a template's discovered placeholder keys to a fill-data map from
    REAL sources only.

    Given the discovered `keys` and the real data handles (person_profile
    `profile` dict, `case_meta` from get_request_case_meta, resolved
    `license_name`, in-session `documents`), classify each key and source it:
      - FORM    → the request's filled SIAP Formulir Isian (`overrides`) — WINS
        over every other source (see below)
      - PROFILE → person_profile.properties (deterministic)
      - CASE    → case meta / licence name (deterministic)
      - VISION  → ONE best-effort Gemini Vision pass over the docs (`run_vision`)
      - BLANK   → never sourced (SIAP-issued or unresolvable)

    `overrides` is the request's `ptsp.form_value.properties` (the filled
    Formulir Isian). For each discovered key, if `overrides` holds a NON-BLANK
    value, that value is used and the key is classed SOURCE_FORM — and NO other
    source (profile/case/Vision) runs for that key. This makes a draft reflect
    EXACTLY what the SIAP form holds (the source of truth), with no redundant /
    divergent re-read. A key ABSENT from `overrides` (or present but blank/
    whitespace) still falls through to profile/case/Vision/blank exactly as
    before — a blank override never overrides a real source with emptiness and
    never fabricates a value.

    Returns `(data, classes)`:
      * `data`   — {key: value} with ONLY the keys that resolved to a NON-EMPTY
        real value. Feed straight to fill_docx_placeholders; every absent key
        renders as the blank fill line.
      * `classes` — {key: source_class} for EVERY discovered key (for the
        officer-facing filled-vs-blank note). A key present in `classes` but
        absent from `data` is a blank the officer completes.

    NEVER raises — a source miss just yields a blank for that key. No value is
    ever fabricated: the resolver adds a key to `data` only after reading it from
    a real source (the form, the profile, the case, or the document).
    """
    classes: dict[str, str] = {}
    data: dict[str, str] = {}
    vision_keys: list[str] = []

    for key in keys:
        # FORM override wins first — a NON-BLANK value from the filled SIAP
        # Formulir Isian short-circuits every other source (and any Vision pass)
        # for this key. A blank/whitespace override is ignored (falls through).
        override = _override_value(key, overrides)
        if override:
            classes[key] = SOURCE_FORM
            data[key] = override
            continue

        cls = classify_placeholder_key(key)
        classes[key] = cls
        if cls == SOURCE_PROFILE:
            val = _resolve_profile_value(key, profile or {})
            if val:
                data[key] = val
        elif cls == SOURCE_CASE:
            val = _resolve_case_value(key, case_meta, license_name)
            if val:
                data[key] = val
        elif cls == SOURCE_VISION:
            vision_keys.append(key)
        # SOURCE_BLANK: deliberately not sourced.

    if run_vision and vision_keys:
        vision = await _extract_vision_fields(vision_keys, documents)
        for key, val in vision.items():
            if val:
                data[key] = val

    logger.info(
        "resolve_fill_data | keys=%d form=%d profile=%d case=%d vision_req=%d filled=%d",
        len(keys),
        sum(1 for c in classes.values() if c == SOURCE_FORM),
        sum(1 for c in classes.values() if c == SOURCE_PROFILE),
        sum(1 for c in classes.values() if c == SOURCE_CASE),
        len(vision_keys),
        len(data),
    )
    return data, classes


async def render_output_docx(
    *,
    internal_filename: Optional[str],
    profile: Optional[dict] = None,
    case_meta: Optional[dict] = None,
    license_name: Optional[str] = None,
    documents: Optional[dict] = None,
    ticket: Optional[str] = None,
    out_prefix: str = "Dokumen",
    overrides: Optional[dict] = None,
) -> Optional[tuple[bytes, str, dict[str, str], dict[str, str]]]:
    """Generic output-template renderer for ANY licence.

    Fetch the fillable `.docx` at `internal_filename`, DISCOVER its `[data.KEY]`
    keys, resolve them from REAL sources (form_value / profile / case / Vision),
    fill, and return `(docx_bytes, filename, data, classes)` where `data` is the
    map that was filled and `classes` is the per-key source classification (both
    power the officer's filled-vs-blank note). Returns None on any miss (disabled,
    no template, not a .docx, 404, python-docx absent, fill error) so the caller
    degrades gracefully. Never raises.

    `overrides` is the request's filled SIAP Formulir Isian
    (`ptsp.form_value.properties`). When supplied, a NON-BLANK value there WINS
    for its key (classed SOURCE_FORM) and no other source runs for that key — so
    the draft reflects exactly what the SIAP form holds. See resolve_fill_data.
    """
    if not _ENABLED:
        return None
    if not internal_filename or not str(internal_filename).lower().endswith(".docx"):
        logger.info("render_output_docx: template not a fillable .docx | file=%s",
                    internal_filename)
        return None

    raw = await fetch_template_bytes(internal_filename)
    if not raw:
        return None

    try:
        keys = discover_placeholder_keys(raw)
    except Exception:
        logger.exception("render_output_docx: placeholder discovery failed")
        return None

    data, classes = await resolve_fill_data(
        keys,
        profile=profile,
        case_meta=case_meta,
        license_name=license_name,
        documents=documents,
        overrides=overrides,
    )

    try:
        filled = fill_docx_placeholders(raw, data)
    except Exception:
        logger.exception("render_output_docx: fill failed")
        return None

    safe_ticket = re.sub(r"\W", "", str(ticket or "")) or "draft"
    safe_prefix = re.sub(r"\W+", "_", str(out_prefix or "Dokumen")).strip("_") or "Dokumen"
    filename = f"{safe_prefix}_{safe_ticket}.docx"
    return filled, filename, data, classes


# ---------------------------------------------------------------------------
# Hand-built DOCX for the legacy .doc Surat Permohonan PPKP (SIAP stores it as
# .doc, which python-docx cannot open and we refuse to ship LibreOffice for).
# Layout mirrors the official template `SURAT PERMOHONAN PPKP.doc` verbatim.
# ---------------------------------------------------------------------------
_PERMOHONAN_FIELD_ORDER = (
    ("Nama", "applicant_name"),
    ("Jabatan", "jabatan"),
    ("Alamat", "alamat"),
    ("Nomer HP", "phone"),
    ("NIK", "nik"),
    ("Nama Kapal", None),          # vessel fields: blank for the citizen
    ("Range GT", None),
    ("Bahan Kapal", None),
    ("Alat Penangkap Ikan", None),
    ("Nama Tukang dan Alamat Galangan", None),
)


def build_surat_permohonan_ppkp_docx(fields: dict[str, Any]) -> bytes:
    """Build the official SIAP Surat Permohonan PPKP as an editable .docx,
    identity fields pre-filled, vessel fields left blank."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Times New Roman"
    normal.size = Pt(12)

    def _p(text="", *, align=None, bold=False, italic=False, size=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
            if size:
                r.font.size = Pt(size)
        return p

    # Date (right-aligned)
    _p(f"Semarang, {_today_id()}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p()
    _p("Kepada Yth.")
    _p("Kepala Dinas Kelautan dan Perikanan Provinsi Jawa Tengah")
    _p("di –")
    _p("\tTEMPAT")
    _p()
    _p("Dengan hormat,")
    _p("Saya yang bertanda tangan di bawah ini, mengajukan Permohonan Persetujuan "
       "Pengadaan Kapal Perikanan (PPKP) Pembangunan / Modifikasi *) sebagai berikut :")

    # Field block as a borderless table: [label | : | value]
    table = doc.add_table(rows=len(_PERMOHONAN_FIELD_ORDER), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for i, (label, field_key) in enumerate(_PERMOHONAN_FIELD_ORDER):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = ":"
        value = _identity_value(field_key, fields) if field_key else None
        cells[2].text = value or ""
        # Make the filled identity values visually distinct (bold), like a form.
        if value:
            for r in cells[2].paragraphs[0].runs:
                r.bold = True

    _p()
    _p("Demikian surat permohonan ini saya sampaikan, atas perhatiannya kami "
       "ucapkan terimakasih.")
    _p()
    _p("Hormat Saya,", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p()
    _p("(Meterai Rp10.000 + Tanda Tangan)", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, size=9)
    _p()
    name = _identity_value("applicant_name", fields) or ""
    _p(name, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    _p("Pemilik Kapal", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p("*) Coret Yang Tidak Perlu", italic=True, size=9)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# doc_type -> hand-built DOCX builder (used when there is no fillable .docx).
_BUILDERS = {
    "surat_permohonan": build_surat_permohonan_ppkp_docx,
}

# Nice download filenames per doc_type.
_OUT_NAMES = {
    "surat_permohonan": "Surat_Permohonan_PPKP.docx",
    "pakta_integritas": "Pakta_Integritas_PPKP.docx",
    "surat_pesanan": "Surat_Pesanan_PPKP.docx",
}


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------
async def render_official_docx(
    license_id: Optional[int], doc_type: str, fields: dict[str, Any]
) -> Optional[tuple[bytes, str]]:
    """Return (docx_bytes, download_filename) for the official SIAP template, or
    None so the caller falls back to the existing fpdf2 generator.

    Order: a synced fillable .docx wins (dynamic, always current); otherwise a
    hand-built DOCX (the legacy .doc case); otherwise None (fallback).
    """
    if not _ENABLED:
        return None

    out_name = _OUT_NAMES.get(doc_type, f"{doc_type}.docx")

    internal: Optional[str] = None
    if license_id is not None:
        try:
            from services.siap_tools import siap_get_templates
            info = await siap_get_templates(int(license_id))
            internal = (info.get("by_doctype") or {}).get(doc_type)
        except Exception:
            logger.exception("siap_get_templates lookup failed | license=%s doc=%s",
                             license_id, doc_type)

    # 1) A real .docx on Beta → fetch + fill (dynamic; reflects SIAP edits).
    if internal and internal.lower().endswith(".docx"):
        raw = await fetch_template_bytes(internal)
        if raw:
            try:
                return fill_docx_identity(raw, fields), out_name
            except Exception:
                logger.exception("Fill of SIAP .docx failed | doc=%s file=%s",
                                 doc_type, internal)

    # 2) Legacy .doc / not-synced / fill error → hand-built DOCX if we have one.
    builder = _BUILDERS.get(doc_type)
    if builder:
        try:
            return builder(fields), out_name
        except Exception:
            logger.exception("Hand-built DOCX failed | doc=%s", doc_type)

    # 3) Nothing usable → fall back to fpdf2.
    return None


# ===========================================================================
# SK / Surat Persetujuan draft — the OFFICER deputy path.
#
# When the officer says "draftkan SK" / "buat surat persetujuan", BIMA fetches
# SIAP's own LICENSE-RECOMMEND output template (the Surat PKPP), fills the
# `[data.KEY]` placeholders it can, and hands back the editable .docx for the
# officer to finish + upload. This is a DRAFT aid, not a SIAP write — SIAP
# issues the SK ber-TTE at final approval. Best-effort throughout: any miss
# (no license_id, template 404, Vision fails, python-docx absent) → None, so
# the caller degrades to an honest message instead of a dead-end.
#
# The 16 PKPP keys, and where each comes from:
#   identity  (deterministic, from the case session):
#     nama_pemohon ← applicant name · alamat ← KTP address ·
#     jenis_pengadaan ← parsed from the licence name (Pembangunan/Modifikasi)
#   vessel/permit (ONE best-effort Gemini Vision pass over the submission docs):
#     nama_kapal, gt, bahan, thn_bangun, alat, galangan, no_siup, tgl_siup,
#     tgl_srt_permohonan, tipe, perihal_surat_perm
#   assigned by SIAP/officer (always left blank):
#     no_ppkp, rekom_tgl_penetapan
# ===========================================================================

# The 16 placeholder keys the PKPP template fills, split by provenance.
_SK_IDENTITY_KEYS = ("nama_pemohon", "alamat", "jenis_pengadaan")
_SK_VESSEL_KEYS = (
    "nama_kapal", "gt", "bahan", "thn_bangun", "alat", "galangan",
    "no_siup", "tgl_siup", "tgl_srt_permohonan", "tipe", "perihal_surat_perm",
)
_SK_OFFICIAL_KEYS = ("no_ppkp", "rekom_tgl_penetapan")  # SIAP/officer assign

# Gemini Vision schema for the vessel/permit fields — pulled from the citizen's
# submission documents (surat permohonan, SIUP, spec sheet). Every field is a
# string; empty when not visible (the officer completes the blanks).
_SK_VESSEL_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "nama_kapal": {"type": "string", "description": "Nama kapal perikanan."},
        "gt": {"type": "string", "description": "Ukuran/tonase kapal (Gross Tonnage), angka saja bila ada."},
        "bahan": {"type": "string", "description": "Bahan kasko kapal (Kayu/Fiberglass/Besi/Baja)."},
        "thn_bangun": {"type": "string", "description": "Tahun pembuatan/pembangunan kapal."},
        "alat": {"type": "string", "description": "Alat penangkap ikan yang digunakan."},
        "galangan": {"type": "string", "description": "Nama/alamat galangan atau tukang pembuat kapal."},
        "no_siup": {"type": "string", "description": "Nomor SIUP (Surat Izin Usaha Perikanan)."},
        "tgl_siup": {"type": "string", "description": "Tanggal terbit SIUP."},
        "tgl_srt_permohonan": {"type": "string", "description": "Tanggal surat permohonan pemohon."},
        "tipe": {"type": "string", "description": "Tipe/jenis kapal."},
        "perihal_surat_perm": {"type": "string", "description": "Perihal surat permohonan."},
    },
    "required": [],
}

_SK_VESSEL_PROMPT = (
    "Anda asisten petugas perizinan kelautan. Dari dokumen permohonan PPKP "
    "terlampir (surat permohonan, SIUP, spesifikasi kapal), ambil data kapal "
    "dan surat yang benar-benar TERTULIS. Untuk setiap field yang tidak "
    "terlihat, kembalikan string kosong. JANGAN mengarang — hanya yang terbaca."
)


def _jenis_pengadaan_from_license(license_name: Optional[str]) -> str:
    """Parse Pembangunan / Modifikasi from the licence name. Returns '' when
    neither word is present (officer fills the blank)."""
    low = (license_name or "").lower()
    if "modifikasi" in low:
        return "Modifikasi"
    if "pembangunan" in low or "pembuatan" in low:
        return "Pembangunan"
    return ""


async def _extract_vessel_fields(documents: Optional[dict]) -> dict[str, str]:
    """Run ONE best-effort Gemini Vision pass over the submission docs to pull
    the vessel/permit fields. Prefers the surat-permohonan/SIUP doc when it can
    identify one; otherwise reads the first doc that carries bytes. Returns a
    {key: value} dict (only non-empty values), or {} when Vision is
    unconfigured / fails / there are no bytes. NEVER raises. Bytes never logged.

    `documents` is the copilot `_doc_context` shape:
      {file_id: {"filename", "mime_type", "content" (bytes), "claimed_type",
                 "detected_type"}}
    """
    if not documents or not isinstance(documents, dict):
        return {}
    try:
        from services.gemini_vision import extract_structured, is_configured
    except Exception:  # pragma: no cover — vision module import guard
        return {}
    if not is_configured():
        logger.info("SK vessel extract: Gemini Vision not configured")
        return {}

    # Pick the most permit-bearing doc: prefer one read/claimed as a
    # surat-permohonan or SIUP, else the first doc with bytes.
    def _label(d: dict) -> str:
        return " ".join(
            str(d.get(k) or "") for k in ("detected_type", "claimed_type", "filename")
        ).lower()

    docs_with_bytes = [d for d in documents.values() if d.get("content")]
    if not docs_with_bytes:
        return {}
    preferred = next(
        (d for d in docs_with_bytes
         if any(w in _label(d) for w in ("permohonan", "siup", "kapal", "izin usaha"))),
        docs_with_bytes[0],
    )

    try:
        parsed = await extract_structured(
            image_bytes=preferred["content"],
            mime_type=preferred.get("mime_type", "application/octet-stream"),
            prompt=_SK_VESSEL_PROMPT,
            response_schema=_SK_VESSEL_SCHEMA,
        )
    except Exception:  # pragma: no cover — extract_structured already guards
        logger.exception("SK vessel extract raised (degrading to blanks)")
        return {}
    if not parsed or not isinstance(parsed, dict):
        return {}
    out: dict[str, str] = {}
    for key in _SK_VESSEL_KEYS:
        val = str(parsed.get(key) or "").strip()
        if val:
            out[key] = val
    logger.info("SK vessel extract | fields=%d", len(out))
    return out


async def build_sk_data(
    *,
    applicant_name: Optional[str],
    alamat: Optional[str],
    license_name: Optional[str],
    documents: Optional[dict] = None,
) -> dict[str, str]:
    """Assemble the `[data.KEY]` fill map for the PKPP SK from what the case
    session holds. Identity fields fill deterministically; vessel/permit fields
    come from ONE best-effort Vision pass; the SIAP/officer-assigned fields are
    left blank. Never raises — a Vision miss just yields more blanks."""
    data: dict[str, str] = {
        "nama_pemohon": (applicant_name or "").strip(),
        "alamat": (alamat or "").strip(),
        "jenis_pengadaan": _jenis_pengadaan_from_license(license_name),
    }
    vessel = await _extract_vessel_fields(documents)
    data.update(vessel)
    # no_ppkp / rekom_tgl_penetapan intentionally absent → rendered as blanks.
    return data


async def render_sk_docx(
    license_id: Optional[int],
    data: dict[str, Any],
    *,
    ticket: Optional[str] = None,
) -> Optional[tuple[bytes, str]]:
    """Render the PKPP approval letter (SK) from SIAP's LICENSE-RECOMMEND
    template, filled with `data`. Returns (docx_bytes, "SK_PPKP_<ticket>.docx")
    or None on ANY miss (disabled, no license_id, no template, 404, fill error,
    python-docx absent) so the caller degrades gracefully. Never raises.
    """
    if not _ENABLED or license_id is None:
        return None

    try:
        from services.siap_tools import siap_get_output_template
        info = await siap_get_output_template(int(license_id))
    except Exception:
        logger.exception("siap_get_output_template lookup failed | license=%s", license_id)
        return None

    if not info.get("found"):
        logger.info("SK render: no output template | license=%s", license_id)
        return None
    internal = info.get("internal_filename")
    if not internal or not str(internal).lower().endswith(".docx"):
        logger.info(
            "SK render: template not a .docx | license=%s file=%s",
            license_id, info.get("file_name"),
        )
        return None

    raw = await fetch_template_bytes(internal)
    if not raw:
        return None

    try:
        filled = fill_docx_placeholders(raw, data)
    except Exception:
        logger.exception("SK render: fill failed | license=%s", license_id)
        return None

    safe_ticket = re.sub(r"\W", "", str(ticket or "")) or "draft"
    return filled, f"SK_PPKP_{safe_ticket}.docx"


# ---------------------------------------------------------------------------
# PDF VIEW copy of the filled SK — so the officer can PREVIEW it on WhatsApp.
#
# Why: WhatsApp/APTANA does NOT reliably render an inline PREVIEW of a .docx
# document message — the officer sees a file card they must download and open in
# Word. During the live rehearsal that produced a ".docx tidak bisa dibuka"
# dead-end. A PDF, by contrast, previews inline in WhatsApp. So alongside the
# EDITABLE .docx we also produce a read-only PDF of the SAME filled content
# (PDF = untuk dibaca, .docx = untuk diedit).
#
# Implementation is PURE-PYTHON — no LibreOffice/soffice in the image. We read
# the just-filled .docx back with python-docx (paragraphs + table cells, in
# document order) and lay the text out with fpdf2. Both libs already ship in
# requirements.txt (fpdf2 for the citizen PPKP path, python-docx for the SIAP
# templates), so this adds NO dependency. It is a plain-text rendering, not a
# pixel-perfect copy — enough for the officer to READ and verify before editing
# the .docx. Best-effort: any miss (either lib absent, parse/render error) →
# None, and the caller simply ships the .docx alone with the honest wording.
# ---------------------------------------------------------------------------

# fpdf2's core fonts are latin-1 only; the fill line uses '…' (U+2026) and SK
# prose may carry other non-latin-1 glyphs. Map the ones we emit, then drop the
# rest, so a stray character can never crash the render.
_PDF_CHAR_MAP = {
    "…": "...", "–": "-", "—": "-", "’": "'", "‘": "'",
    "“": '"', "”": '"', " ": " ",
}


def _pdf_safe_text(text: str) -> str:
    """Make `text` safe for fpdf2's latin-1 core fonts: swap the smart-quote /
    ellipsis / dash glyphs we know we emit, then strip anything still outside
    latin-1 so encoding can never raise."""
    for bad, good in _PDF_CHAR_MAP.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "replace").decode("latin-1")


def _iter_docx_lines(docx_bytes: bytes) -> list[str]:
    """Extract the filled document's text as ordered lines (paragraphs, then each
    table row joined by '  '), skipping blank rows. Used to lay the SK out as a
    readable PDF view copy."""
    from docx import Document  # lazy: only when producing the PDF view

    doc = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(para.text.rstrip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            joined = "  ".join(c for c in cells if c)
            if joined:
                lines.append(joined)
    return lines


def render_pdf_view_from_docx(
    docx_bytes: bytes, *, ticket: Optional[str] = None
) -> Optional[tuple[bytes, str]]:
    """Render a READ-ONLY PDF preview of an already-filled SK .docx, so the
    officer can view it inline on WhatsApp (which does not preview .docx).

    Returns (pdf_bytes, "SK_PPKP_<ticket>.pdf") or None on ANY miss (fpdf2 or
    python-docx absent, parse/render error). Pure-Python; never raises. The PDF
    carries the SAME filled content as the .docx but is not editable — the .docx
    remains the source of truth for the officer to complete.
    """
    try:
        from fpdf import FPDF
    except Exception:  # pragma: no cover — dependency guard
        logger.info("SK PDF view: fpdf2 not available, shipping .docx only")
        return None

    try:
        lines = _iter_docx_lines(docx_bytes)
    except Exception:
        logger.exception("SK PDF view: could not read filled .docx")
        return None

    try:
        pdf = FPDF(format="A4", unit="mm")
        # Explicit margins so the effective text width is a known positive value
        # (fpdf2 raises "not enough horizontal space" if a multi_cell width can't
        # fit a character — passing the computed width defensively avoids it).
        margin = 15.0
        pdf.set_margins(margin, margin, margin)
        pdf.set_auto_page_break(auto=True, margin=margin)
        pdf.add_page()
        eff_w = pdf.w - 2 * margin  # printable width in mm
        # A slim header so the officer sees at a glance this is the read copy.
        pdf.set_font("Helvetica", style="B", size=9)
        pdf.set_text_color(120, 120, 120)
        pdf.multi_cell(
            eff_w, 5,
            _pdf_safe_text(
                "SALINAN UNTUK DIBACA (Draf SK) - versi .docx yang bisa diedit "
                "dikirim terpisah."
            ),
        )
        pdf.ln(2)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", size=11)
        for line in lines:
            safe = _pdf_safe_text(line)
            if safe.strip():
                pdf.multi_cell(eff_w, 6, safe)
            else:
                pdf.ln(3)  # preserve blank-line spacing
        out = pdf.output()
        pdf_bytes = bytes(out)
    except Exception:
        logger.exception("SK PDF view: fpdf2 render failed")
        return None

    safe_ticket = re.sub(r"\W", "", str(ticket or "")) or "draft"
    return pdf_bytes, f"SK_PPKP_{safe_ticket}.pdf"
