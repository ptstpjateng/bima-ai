"""Official SIAP template rendering (services/siap_templates.py).

Covers the label matcher, the hand-built Surat Permohonan PPKP DOCX, and the
in-place identity fill. Skips automatically where python-docx / httpx are not
installed (the bare CI shell) — it runs in the ai-engine container and any venv
with the deps. The dynamic fetch (render_official_docx over the SIAP DB +
Beta storage) is verified live against Beta post-deploy.
"""
from __future__ import annotations

import io
import sys
import types
import unittest
from pathlib import Path
from unittest.mock import patch

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def _ensure_stub(name: str, attrs: dict | None = None) -> None:
    """Install a thin stub module ONLY when the real one is absent (bare env).
    Never overwrites a real module already imported — so a sibling suite that
    needs the real services.siap_tools (which pulls asyncpg/dotenv) is not
    poisoned. We stub only the leaf infra deps, never a services.* module."""
    if name in sys.modules:
        return
    try:
        __import__(name)
        return
    except ImportError:
        pass
    mod = types.ModuleType(name)
    for attr, value in (attrs or {}).items():
        setattr(mod, attr, value)
    sys.modules[name] = mod


# render_sk_docx lazily imports services.siap_tools, which imports services.
# siap_db → asyncpg + dotenv. Stub those leaf infra deps so the SK-render tests
# can drive the (patched) siap_get_output_template without a live DB driver.
_ensure_stub("asyncpg")
_ensure_stub("dotenv", {"load_dotenv": lambda *a, **k: None})

try:
    import docx  # noqa: F401  (fill/build need it; lazy inside the module)
    from docx import Document
    from services import siap_templates as st  # imports httpx at top
    _DEPS_OK = True
except Exception:  # pragma: no cover — bare env without python-docx/httpx
    _DEPS_OK = False

_FIELDS = {
    "applicant_name": "CASMO",
    "nik": "3374012345678901",
    "phone": "628123456789",
    "alamat": "Jl. Laut No. 1, Tegal",
}


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestLabelMatcher(unittest.TestCase):
    def test_identity_labels_match(self):
        self.assertEqual(st._match_identity_field("Nama"), "applicant_name")
        self.assertEqual(st._match_identity_field("NIK"), "nik")
        self.assertEqual(st._match_identity_field("Alamat"), "alamat")
        self.assertEqual(st._match_identity_field("Jabatan"), "jabatan")
        self.assertEqual(st._match_identity_field("No HP"), "phone")
        self.assertEqual(st._match_identity_field("Nomer HP"), "phone")

    def test_whole_label_only_no_substring_bleed(self):
        # These must NOT be treated as identity fields (would misfill).
        self.assertIsNone(st._match_identity_field("Nama Kapal"))
        self.assertIsNone(st._match_identity_field("Nama lembaga/institusi"))
        self.assertIsNone(st._match_identity_field("Nama Penyedia/Galangan"))
        self.assertIsNone(st._match_identity_field("Jabatan dalam lembaga/institusi"))


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestSuratPermohonanBuilder(unittest.TestCase):
    def _doc(self):
        return Document(io.BytesIO(st.build_surat_permohonan_ppkp_docx(_FIELDS)))

    def test_correct_recipient_not_dpmptsp(self):
        paras = "\n".join(p.text for p in self._doc().paragraphs)
        self.assertIn("Dinas Kelautan dan Perikanan Provinsi Jawa Tengah", paras)
        self.assertNotIn("DPMPTSP", paras)
        self.assertIn("Coret Yang Tidak Perlu", paras)

    def test_identity_filled_vessel_blank(self):
        doc = self._doc()
        rows = {t.rows[i].cells[0].text: t.rows[i].cells[2].text
                for t in doc.tables for i in range(len(t.rows))}
        self.assertEqual(rows.get("Nama"), "CASMO")
        self.assertEqual(rows.get("NIK"), "3374012345678901")
        self.assertEqual(rows.get("Nomer HP"), "628123456789")
        self.assertEqual(rows.get("Alamat"), "Jl. Laut No. 1, Tegal")
        self.assertEqual(rows.get("Jabatan"), "Pemilik Kapal")
        for vessel in ("Nama Kapal", "Range GT", "Bahan Kapal",
                       "Alat Penangkap Ikan", "Nama Tukang dan Alamat Galangan"):
            self.assertEqual(rows.get(vessel), "", f"{vessel} must be blank")

    def test_output_is_valid_docx(self):
        data = st.build_surat_permohonan_ppkp_docx(_FIELDS)
        self.assertEqual(data[:2], b"PK")  # DOCX = zip container


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestFillDocxIdentity(unittest.TestCase):
    def _template(self, labels):
        d = Document()
        for lbl in labels:
            d.add_paragraph(f"{lbl}\t: ")
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    def _lines(self, docx_bytes):
        fd = Document(io.BytesIO(docx_bytes))
        out = {}
        for p in st._iter_block_paragraphs(fd):
            if ":" in p.text:
                k, _, v = p.text.partition(":")
                out[k.strip()] = v.strip()
        return out

    def test_fills_identity_leaves_vessel(self):
        raw = self._template(["Nama", "Alamat", "No HP", "NIK", "Nama Kapal", "Range GT"])
        lines = self._lines(st.fill_docx_identity(raw, _FIELDS))
        self.assertEqual(lines["Nama"], "CASMO")
        self.assertEqual(lines["Alamat"], "Jl. Laut No. 1, Tegal")
        self.assertEqual(lines["No HP"], "628123456789")
        self.assertEqual(lines["NIK"], "3374012345678901")
        self.assertEqual(lines["Nama Kapal"], "")   # vessel untouched
        self.assertEqual(lines["Range GT"], "")

    def test_does_not_overwrite_prefilled_slot(self):
        d = Document()
        d.add_paragraph("Nama\t: Sudah Terisi")
        buf = io.BytesIO(); d.save(buf)
        lines = self._lines(st.fill_docx_identity(buf.getvalue(), _FIELDS))
        self.assertEqual(lines["Nama"], "Sudah Terisi")  # not clobbered

    def test_missing_field_left_blank(self):
        raw = self._template(["Nama", "Alamat"])
        # No alamat in fields → Alamat stays blank; Nama fills.
        lines = self._lines(st.fill_docx_identity(raw, {"applicant_name": "CASMO"}))
        self.assertEqual(lines["Nama"], "CASMO")
        self.assertEqual(lines["Alamat"], "")


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestFillDocxPlaceholders(unittest.TestCase):
    """The `[data.KEY]` fill for the OFFICER SK output template — run-aware,
    blank-fills unmapped keys, never leaves a raw token behind."""

    def _para_from_runs(self, doc, run_texts):
        """Build a paragraph whose text is SPLIT across the given runs — this is
        exactly how Word shatters a `[data.KEY]` token across runs, the case a
        naive per-run replace would miss."""
        p = doc.add_paragraph()
        for t in run_texts:
            p.add_run(t)
        return p

    def _template_with_split_token(self):
        """A template where `[data.no_siup]` is split across three runs (like the
        real Surat PKPP), plus an intact `[data.gt]` and a to-be-blanked
        `[data.tgl_siup]`."""
        d = Document()
        # no_siup shattered: "[", "data.no_siup", "]" — plus a trailing token.
        self._para_from_runs(
            d, ["SIUP Nomor : ", "[", "data.no_siup", "]",
                " Tanggal ", "[data.tgl_siup]"],
        )
        # gt intact in a single run.
        self._para_from_runs(d, ["Ukuran GT: ", "[data.gt]"])
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    def _text(self, docx_bytes):
        fd = Document(io.BytesIO(docx_bytes))
        return "\n".join(p.text for p in st._iter_block_paragraphs(fd))

    def test_fills_run_split_token(self):
        raw = self._template_with_split_token()
        out = st.fill_docx_placeholders(raw, {"no_siup": "SIUP-123", "gt": "5"})
        text = self._text(out)
        # The shattered token filled from the JOINED paragraph text.
        self.assertIn("SIUP-123", text)
        self.assertIn("5", text)
        # No raw token survives anywhere.
        self.assertNotIn("[data.", text)

    def test_unmapped_and_empty_keys_become_blank_fill(self):
        raw = self._template_with_split_token()
        # tgl_siup is absent; no_siup empty → both become the blank fill line,
        # NOT the raw "[data.KEY]".
        out = st.fill_docx_placeholders(raw, {"no_siup": "", "gt": "5"})
        text = self._text(out)
        self.assertNotIn("[data.", text)
        self.assertIn(st._BLANK_FILL, text)

    def test_multiple_tokens_one_paragraph(self):
        d = Document()
        self._para_from_runs(
            d, ["Kapal ", "[data.nama_kapal]", " bahan ", "[data.bahan]"])
        buf = io.BytesIO(); d.save(buf)
        out = st.fill_docx_placeholders(
            buf.getvalue(), {"nama_kapal": "KM BAHARI", "bahan": "Kayu"})
        text = self._text(out)
        self.assertIn("KM BAHARI", text)
        self.assertIn("Kayu", text)
        self.assertNotIn("[data.", text)

    def test_output_is_valid_docx(self):
        raw = self._template_with_split_token()
        out = st.fill_docx_placeholders(raw, {"gt": "5"})
        self.assertEqual(out[:2], b"PK")

    def test_fill_in_table_cell(self):
        # Tokens inside a table cell are reached by _iter_block_paragraphs too.
        d = Document()
        t = d.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        cell.paragraphs[0].add_run("Tipe: ")
        cell.paragraphs[0].add_run("[data.tipe]")
        buf = io.BytesIO(); d.save(buf)
        out = st.fill_docx_placeholders(buf.getvalue(), {"tipe": "Fiber"})
        self.assertIn("Fiber", self._text(out))
        self.assertNotIn("[data.", self._text(out))


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestBuildSkData(unittest.IsolatedAsyncioTestCase):
    """build_sk_data — identity deterministic, jenis_pengadaan parsed from the
    licence name, vessel fields from a best-effort (here: absent) Vision pass."""

    async def test_identity_and_jenis_pengadaan_pembangunan(self):
        data = await st.build_sk_data(
            applicant_name="CASMO", alamat="Jl. Laut No. 1",
            license_name="Persetujuan Pengadaan Kapal (Pembangunan)",
            documents={},
        )
        self.assertEqual(data["nama_pemohon"], "CASMO")
        self.assertEqual(data["alamat"], "Jl. Laut No. 1")
        self.assertEqual(data["jenis_pengadaan"], "Pembangunan")

    async def test_jenis_pengadaan_modifikasi(self):
        data = await st.build_sk_data(
            applicant_name="A", alamat="B",
            license_name="Izin ... Modifikasi Kapal", documents=None)
        self.assertEqual(data["jenis_pengadaan"], "Modifikasi")

    async def test_jenis_pengadaan_blank_when_absent(self):
        data = await st.build_sk_data(
            applicant_name="A", alamat="B", license_name="Izin Lain", documents=None)
        self.assertEqual(data["jenis_pengadaan"], "")

    async def test_vessel_fields_from_vision(self):
        # One in-session doc with bytes → build_sk_data runs ONE Vision pass and
        # merges the non-empty vessel fields. Stub gemini_vision.
        mod = types.ModuleType("services.gemini_vision")

        async def _extract(*, image_bytes, mime_type, prompt, response_schema):
            return {"nama_kapal": "KM BAHARI", "gt": "5", "bahan": "",
                    "no_siup": "SIUP-9"}

        mod.extract_structured = _extract
        mod.is_configured = lambda: True
        sys.modules["services.gemini_vision"] = mod
        try:
            data = await st.build_sk_data(
                applicant_name="CASMO", alamat="Jl X",
                license_name="Pengadaan Kapal (Pembangunan)",
                documents={"doc-1": {"filename": "permohonan.pdf",
                                     "mime_type": "application/pdf",
                                     "content": b"pdfbytes",
                                     "claimed_type": "Surat_Permohonan"}},
            )
        finally:
            sys.modules.pop("services.gemini_vision", None)
        self.assertEqual(data["nama_kapal"], "KM BAHARI")
        self.assertEqual(data["gt"], "5")
        self.assertEqual(data["no_siup"], "SIUP-9")
        # Empty vessel value not carried; identity still present.
        self.assertNotIn("bahan", data)
        self.assertEqual(data["nama_pemohon"], "CASMO")

    async def test_vision_unconfigured_leaves_vessel_blank(self):
        mod = types.ModuleType("services.gemini_vision")
        mod.extract_structured = None
        mod.is_configured = lambda: False
        sys.modules["services.gemini_vision"] = mod
        try:
            data = await st.build_sk_data(
                applicant_name="CASMO", alamat="Jl X",
                license_name="Pengadaan (Pembangunan)",
                documents={"doc-1": {"content": b"x", "mime_type": "application/pdf"}})
        finally:
            sys.modules.pop("services.gemini_vision", None)
        # No vessel keys added; identity intact.
        self.assertNotIn("nama_kapal", data)
        self.assertEqual(data["nama_pemohon"], "CASMO")


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestRenderSkDocx(unittest.IsolatedAsyncioTestCase):
    """render_sk_docx — fetch LICENSE-RECOMMEND template, fill, name the file;
    None on every miss so the caller degrades."""

    def _template_bytes(self):
        d = Document()
        p = d.add_paragraph()
        for t in ["Pemohon: ", "[", "data.nama_pemohon", "]"]:
            p.add_run(t)
        d.add_paragraph("No PPKP: [data.no_ppkp]")
        buf = io.BytesIO(); d.save(buf)
        return buf.getvalue()

    def _patch(self, *, out_template, fetch_bytes):
        """Patch siap_get_output_template (in siap_tools, imported lazily by
        render_sk_docx) and fetch_template_bytes (on the module)."""
        import services.siap_tools as stools

        async def _out(lid):
            return out_template

        async def _fetch(internal):
            return fetch_bytes

        self._old_out = getattr(stools, "siap_get_output_template", None)
        self._old_fetch = st.fetch_template_bytes
        stools.siap_get_output_template = _out
        st.fetch_template_bytes = _fetch

    def _unpatch(self):
        import services.siap_tools as stools
        if self._old_out is not None:
            stools.siap_get_output_template = self._old_out
        st.fetch_template_bytes = self._old_fetch

    async def test_renders_named_docx_from_template(self):
        self._patch(
            out_template={"found": True, "license_id": 459,
                          "internal_filename": "master/ULID.docx"},
            fetch_bytes=self._template_bytes(),
        )
        try:
            res = await st.render_sk_docx(
                459, {"nama_pemohon": "CASMO"}, ticket="000123456")
        finally:
            self._unpatch()
        self.assertIsNotNone(res)
        data, name = res
        self.assertEqual(name, "SK_PPKP_000123456.docx")
        self.assertEqual(data[:2], b"PK")
        text = "\n".join(p.text for p in st._iter_block_paragraphs(
            Document(io.BytesIO(data))))
        self.assertIn("CASMO", text)
        self.assertNotIn("[data.", text)  # no raw token left

    async def test_none_when_no_license_id(self):
        res = await st.render_sk_docx(None, {}, ticket="1")
        self.assertIsNone(res)

    async def test_none_when_template_not_found(self):
        self._patch(out_template={"found": False}, fetch_bytes=None)
        try:
            res = await st.render_sk_docx(459, {}, ticket="1")
        finally:
            self._unpatch()
        self.assertIsNone(res)

    async def test_none_when_template_is_legacy_doc(self):
        self._patch(
            out_template={"found": True, "internal_filename": "master/OLD.doc"},
            fetch_bytes=b"whatever",
        )
        try:
            res = await st.render_sk_docx(459, {}, ticket="1")
        finally:
            self._unpatch()
        self.assertIsNone(res)

    async def test_none_when_fetch_misses(self):
        self._patch(
            out_template={"found": True, "internal_filename": "master/ULID.docx"},
            fetch_bytes=None,   # 404 / not synced
        )
        try:
            res = await st.render_sk_docx(459, {}, ticket="1")
        finally:
            self._unpatch()
        self.assertIsNone(res)


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestFetchSubmissionFileBytes(unittest.IsolatedAsyncioTestCase):
    """fetch_submission_file_bytes — officer copilot pulls a citizen upload from
    Beta storage, reusing the hardened fetch_template_bytes GET. Verifies the
    verbatim-then-prefixed candidate order and graceful None on total miss."""

    async def test_verbatim_hit_returns_bytes_without_prefix_attempt(self):
        seen: list[str] = []

        async def fake_fetch(path):
            seen.append(path)
            return b"PDFDATA"  # first candidate hits

        with patch.object(st, "fetch_template_bytes", new=fake_fetch):
            out = await st.fetch_submission_file_bytes("berkas/2026/ktp.pdf")
        self.assertEqual(out, b"PDFDATA")
        # Verbatim tried first; a hit short-circuits before the prefixed one.
        self.assertEqual(seen, ["berkas/2026/ktp.pdf"])

    async def test_falls_back_to_public_prefix_on_first_miss(self):
        seen: list[str] = []

        async def fake_fetch(path):
            seen.append(path)
            return b"OK" if path.startswith("app/public/") else None

        with patch.object(st, "fetch_template_bytes", new=fake_fetch):
            out = await st.fetch_submission_file_bytes("berkas/x.pdf")
        self.assertEqual(out, b"OK")
        self.assertEqual(seen, ["berkas/x.pdf", "app/public/berkas/x.pdf"])

    async def test_none_when_all_candidates_miss(self):
        async def fake_fetch(path):
            return None

        with patch.object(st, "fetch_template_bytes", new=fake_fetch):
            out = await st.fetch_submission_file_bytes("berkas/gone.pdf")
        self.assertIsNone(out)

    async def test_empty_ref_returns_none_without_fetch(self):
        called = False

        async def fake_fetch(path):
            nonlocal called
            called = True
            return b"x"

        with patch.object(st, "fetch_template_bytes", new=fake_fetch):
            out = await st.fetch_submission_file_bytes("")
        self.assertIsNone(out)
        self.assertFalse(called)


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestPlaceholderDiscovery(unittest.TestCase):
    """discover_placeholder_keys — run-aware, ordered, de-duplicated."""

    def _doc(self, run_groups):
        d = Document()
        for runs in run_groups:
            p = d.add_paragraph()
            for t in runs:
                p.add_run(t)
        buf = io.BytesIO(); d.save(buf)
        return buf.getvalue()

    def test_discovers_intact_and_shattered_tokens_in_order(self):
        raw = self._doc([
            ["Pemohon: ", "[data.nama_pemohon]"],
            ["SIUP ", "[", "data.no_siup", "]", " tgl ", "[data.tgl_siup]"],
            ["Ukuran: ", "[data.gt]"],
        ])
        keys = st.discover_placeholder_keys(raw)
        self.assertEqual(
            keys, ["nama_pemohon", "no_siup", "tgl_siup", "gt"])

    def test_dedupes_repeated_key(self):
        raw = self._doc([["[data.nama_pemohon] ... [data.nama_pemohon]"]])
        self.assertEqual(st.discover_placeholder_keys(raw), ["nama_pemohon"])

    def test_no_tokens_is_empty(self):
        raw = self._doc([["Tidak ada token di sini."]])
        self.assertEqual(st.discover_placeholder_keys(raw), [])


class TestClassifyPlaceholderKey(unittest.TestCase):
    """classify_placeholder_key — the no-hallucination source router. Pure (no
    docx), so it runs even in the bare env."""

    def test_profile_keys(self):
        for k in ("nama_pemohon", "alamat", "nik", "no_hp", "email",
                  "atas_nama_perusahaan", "alamat_perusahaan", "pekerjaan_pemohon"):
            self.assertEqual(st.classify_placeholder_key(k), st.SOURCE_PROFILE, k)

    def test_case_keys(self):
        for k in ("no_tiket", "ticket_num", "nama_izin", "jenis_pengadaan"):
            self.assertEqual(st.classify_placeholder_key(k), st.SOURCE_CASE, k)

    def test_siap_issued_keys_are_blank(self):
        # SK/reg number, penetapan/TTE date, QR/signature → NEVER sourced.
        for k in ("no_ppkp", "no_sk", "no_rekom", "rekom_tgl_penetapan",
                  "tgl_penetapan", "qrcode_signed", "ttd_gbr", "masa_berlaku",
                  "nomor_rekomtek"):
            self.assertEqual(st.classify_placeholder_key(k), st.SOURCE_BLANK, k)

    def test_doc_only_keys_go_to_vision(self):
        for k in ("nama_kapal", "gt", "bahan", "galangan", "luas_tanah",
                  "status_tanah", "batas_tanah"):
            self.assertEqual(st.classify_placeholder_key(k), st.SOURCE_VISION, k)

    def test_unknown_key_falls_to_vision_not_guess(self):
        # A never-seen doc-only key on a NEW template → Vision (best-effort),
        # which returns "" when not visible → blank. Never fabricated.
        self.assertEqual(
            st.classify_placeholder_key("some_new_field"), st.SOURCE_VISION)

    def test_empty_key_is_blank(self):
        self.assertEqual(st.classify_placeholder_key(""), st.SOURCE_BLANK)


class TestResolveFillData(unittest.IsolatedAsyncioTestCase):
    """resolve_fill_data — the DATA-SOURCE RESOLVER. Profile/case deterministic,
    Vision best-effort, everything unresolved stays OUT of the data map (→ the
    renderer emits a blank fill line). No value is ever fabricated."""

    async def test_profile_and_case_resolve_deterministically(self):
        keys = ["nama_pemohon", "alamat", "nik", "no_tiket", "nama_izin",
                "jenis_pengadaan", "no_ppkp"]
        profile = {
            "full_name": "CASMO", "address": "Jl. Laut No. 1",
            "identity_number": "3374012345678901",
        }
        case = {"ticket": "000123456"}
        data, classes = await st.resolve_fill_data(
            keys, profile=profile, case_meta=case,
            license_name="Pengadaan Kapal (Pembangunan)",
            documents=None, run_vision=False,
        )
        self.assertEqual(data["nama_pemohon"], "CASMO")
        self.assertEqual(data["alamat"], "Jl. Laut No. 1")
        self.assertEqual(data["nik"], "3374012345678901")
        self.assertEqual(data["no_tiket"], "000123456")
        self.assertEqual(data["nama_izin"], "Pengadaan Kapal (Pembangunan)")
        self.assertEqual(data["jenis_pengadaan"], "Pembangunan")
        # SIAP-issued → NOT sourced → absent → renders blank.
        self.assertNotIn("no_ppkp", data)
        self.assertEqual(classes["no_ppkp"], st.SOURCE_BLANK)
        # Every discovered key is classified (for the filled/blank note).
        self.assertEqual(set(classes.keys()), set(keys))

    async def test_missing_profile_field_stays_blank(self):
        # Profile lacks address → alamat is NOT added to data (blank), never
        # guessed from another field.
        data, classes = await st.resolve_fill_data(
            ["nama_pemohon", "alamat"],
            profile={"full_name": "CASMO"}, case_meta={},
            license_name=None, documents=None, run_vision=False,
        )
        self.assertEqual(data["nama_pemohon"], "CASMO")
        self.assertNotIn("alamat", data)
        self.assertEqual(classes["alamat"], st.SOURCE_PROFILE)

    async def test_vision_fills_doc_only_keys(self):
        # A VISION key with a value visible in the doc is filled; one not visible
        # stays blank. Stub gemini_vision.
        mod = types.ModuleType("services.gemini_vision")
        captured = {}

        async def _extract(*, image_bytes, mime_type, prompt, response_schema):
            captured["schema_keys"] = sorted(response_schema["properties"].keys())
            return {"nama_kapal": "KM BAHARI", "gt": ""}

        mod.extract_structured = _extract
        mod.is_configured = lambda: True
        sys.modules["services.gemini_vision"] = mod
        try:
            data, classes = await st.resolve_fill_data(
                ["nama_pemohon", "nama_kapal", "gt"],
                profile={"full_name": "CASMO"}, case_meta={},
                license_name=None,
                documents={"d1": {"content": b"x", "mime_type": "application/pdf"}},
            )
        finally:
            sys.modules.pop("services.gemini_vision", None)
        # Schema built ONLY from the VISION-classified keys (data-driven, no
        # hardcode) — nama_pemohon is profile so it's not in the vision schema.
        self.assertEqual(captured["schema_keys"], ["gt", "nama_kapal"])
        self.assertEqual(data["nama_pemohon"], "CASMO")
        self.assertEqual(data["nama_kapal"], "KM BAHARI")
        self.assertNotIn("gt", data)          # empty vision value → blank
        self.assertEqual(classes["gt"], st.SOURCE_VISION)

    async def test_vision_unconfigured_leaves_doc_keys_blank(self):
        mod = types.ModuleType("services.gemini_vision")
        mod.extract_structured = None
        mod.is_configured = lambda: False
        sys.modules["services.gemini_vision"] = mod
        try:
            data, classes = await st.resolve_fill_data(
                ["nama_kapal", "gt"], profile={}, case_meta={},
                license_name=None,
                documents={"d1": {"content": b"x", "mime_type": "application/pdf"}},
            )
        finally:
            sys.modules.pop("services.gemini_vision", None)
        self.assertEqual(data, {})            # nothing sourced → all blank
        self.assertEqual(classes["nama_kapal"], st.SOURCE_VISION)


@unittest.skipUnless(_DEPS_OK, "python-docx/httpx not installed")
class TestRenderOutputDocx(unittest.IsolatedAsyncioTestCase):
    """render_output_docx — the generic renderer: fetch → discover → resolve →
    fill, returning (bytes, name, data, classes). Verifies an UNKNOWN licence
    template (different key set than PKPP) works with NO code change and that
    unresolved keys render blank."""

    def _template(self, run_groups):
        d = Document()
        for runs in run_groups:
            p = d.add_paragraph()
            for t in runs:
                p.add_run(t)
        buf = io.BytesIO(); d.save(buf)
        return buf.getvalue()

    async def test_fills_unknown_license_from_real_sources(self):
        # A pengairan-style template with a DIFFERENT key set than PKPP.
        raw = self._template([
            ["Nama Pemohon: ", "[data.nama_pemohon]"],
            ["Alamat: ", "[data.alamat_pemohon]"],
            ["Luas Tanah: ", "[", "data.luas_tanah", "]"],
            ["Nomor Rekomtek: ", "[data.nomor_rekomtek]"],  # SIAP-issued → blank
        ])

        async def _fetch(internal):
            return raw

        # Stub Vision so luas_tanah resolves from the doc.
        vmod = types.ModuleType("services.gemini_vision")

        async def _extract(*, image_bytes, mime_type, prompt, response_schema):
            return {"luas_tanah": "250 m2"}

        vmod.extract_structured = _extract
        vmod.is_configured = lambda: True
        sys.modules["services.gemini_vision"] = vmod

        old_fetch = st.fetch_template_bytes
        st.fetch_template_bytes = _fetch
        try:
            res = await st.render_output_docx(
                internal_filename="master/REK.docx",
                profile={"full_name": "SUTRISNO", "address": "Jl. Air No. 9"},
                case_meta={"ticket": "000900900"},
                license_name="Izin Pemanfaatan Air (IPTB)",
                documents={"d1": {"content": b"x", "mime_type": "application/pdf"}},
                ticket="000900900",
                out_prefix="Rekomtek",
            )
        finally:
            st.fetch_template_bytes = old_fetch
            sys.modules.pop("services.gemini_vision", None)

        self.assertIsNotNone(res)
        docx_bytes, name, data, classes = res
        self.assertEqual(name, "Rekomtek_000900900.docx")
        self.assertEqual(docx_bytes[:2], b"PK")
        # Real sources filled; SIAP-issued left blank.
        self.assertEqual(data["nama_pemohon"], "SUTRISNO")
        self.assertEqual(data["alamat_pemohon"], "Jl. Air No. 9")
        self.assertEqual(data["luas_tanah"], "250 m2")
        self.assertNotIn("nomor_rekomtek", data)
        self.assertEqual(classes["nomor_rekomtek"], st.SOURCE_BLANK)
        # The rendered doc carries the values + a blank line for the issued key,
        # and NO raw token survives.
        text = "\n".join(p.text for p in st._iter_block_paragraphs(
            Document(io.BytesIO(docx_bytes))))
        self.assertIn("SUTRISNO", text)
        self.assertIn("250 m2", text)
        self.assertIn(st._BLANK_FILL, text)   # nomor_rekomtek blank line
        self.assertNotIn("[data.", text)

    async def test_none_when_not_a_docx(self):
        res = await st.render_output_docx(internal_filename="master/OLD.doc")
        self.assertIsNone(res)

    async def test_none_when_fetch_misses(self):
        async def _fetch(internal):
            return None
        old = st.fetch_template_bytes
        st.fetch_template_bytes = _fetch
        try:
            res = await st.render_output_docx(internal_filename="master/X.docx")
        finally:
            st.fetch_template_bytes = old
        self.assertIsNone(res)


class TestPhpWordAndDigitKeys(unittest.TestCase):
    """LICENSE-SK templates use PhpWord ${key}, not [data.KEY]; some RECOMMEND
    templates carry digit keys (catatan1). The engine must discover AND fill
    both syntaxes (run-aware) and blank issuance/signature anchors."""

    def _para_from_runs(self, doc, run_texts):
        p = doc.add_paragraph()
        for t in run_texts:
            p.add_run(t)
        return p

    def _sk_template(self):
        d = Document()
        # ${nama_pemohon} run-shattered like Word does; ${no_sk} + ${qrcode_signed}
        # are issuance/signature anchors that must blank; a [data.catatan1] digit
        # key must also fill.
        self._para_from_runs(d, ["Yth. ", "$", "{nama_pemohon", "}", ","])
        self._para_from_runs(d, ["Nomor SK: ", "${no_sk}", "  QR: ", "${qrcode_signed}"])
        self._para_from_runs(d, ["Catatan: ", "[data.catatan1]"])
        buf = io.BytesIO(); d.save(buf); return buf.getvalue()

    def _text(self, docx_bytes):
        fd = Document(io.BytesIO(docx_bytes))
        return "\n".join(p.text for p in st._iter_block_paragraphs(fd))

    def test_discovers_both_syntaxes_and_digit_keys(self):
        keys = st.discover_placeholder_keys(self._sk_template())
        self.assertIn("nama_pemohon", keys)   # ${key}, run-shattered
        self.assertIn("no_sk", keys)
        self.assertIn("qrcode_signed", keys)
        self.assertIn("catatan1", keys)        # [data.KEY] with a digit

    def test_fills_phpword_blanks_issuance_and_signature(self):
        out = st.fill_docx_placeholders(
            self._sk_template(),
            {"nama_pemohon": "CASMO", "catatan1": "lengkap"},
        )
        text = self._text(out)
        self.assertIn("CASMO", text)           # ${nama_pemohon} filled
        self.assertIn("lengkap", text)         # [data.catatan1] filled
        self.assertNotIn("${", text)           # no raw PhpWord token left
        self.assertNotIn("[data.", text)
        # no_sk / qrcode_signed are SIAP-issued/signature → classified BLANK →
        # never filled with a fabricated value, rendered as the fill line.
        self.assertEqual(st.classify_placeholder_key("no_sk"), st.SOURCE_BLANK)
        self.assertEqual(st.classify_placeholder_key("qrcode_signed"), st.SOURCE_BLANK)
        self.assertIn(st._BLANK_FILL, text)


if __name__ == "__main__":
    unittest.main(verbosity=2)
